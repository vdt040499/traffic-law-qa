#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Enhanced extraction system for Vietnamese traffic law documents
Supports multiple document formats and incremental updates
"""

import json
import os
import re
from datetime import datetime
from docx import Document
import hashlib

class VietnameseTrafficLawExtractor:
    """Main extractor class for Vietnamese traffic law documents"""
    
    def __init__(self, config_path=None):
        """Initialize extractor with configuration"""
        self.config = self.load_config(config_path)
        self.patterns = self.setup_patterns()
        
    def load_config(self, config_path):
        """Load extraction configuration"""
        default_config = {
            "supported_formats": [".docx", ".pdf", ".txt", ".json"],
            "output_dir": "data/processed",
            "backup_dir": "data/backups",
            "main_document": "data/raw/legal_documents/nghi_dinh_100_2019.json",
            "update_strategy": "merge",  # merge, replace, append
            "auto_detect_changes": True,
            "create_backups": True
        }
        
        if config_path and os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                user_config = json.load(f)
                default_config.update(user_config)
        
        return default_config
    
    def setup_patterns(self):
        """Setup regex patterns for different document structures"""
        return {
            # Article patterns for different decree formats
            "article_patterns": [
                r"Điều\s+(\d+)\.\s*(.*)",  # Standard: Điều 5. Title
                r"Điều\s+(\d+):\s*(.*)",   # Alternative: Điều 5: Title
                r"ĐIỀU\s+(\d+)\.\s*(.*)",  # Uppercase version
            ],
            
            # Section patterns for penalty ranges
            "section_patterns": [
                r"(\d+)\.\s*Phạt tiền từ\s*([\d.,]+)\s*đồng đến\s*([\d.,]+)\s*đồng.*:",
                r"(\d+)\.\s*Phạt tiền từ\s*([\d.,]+)\s*VNĐ đến\s*([\d.,]+)\s*VNĐ.*:",
                r"Khoản\s+(\d+)\.\s*Phạt tiền từ\s*([\d.,]+)\s*đến\s*([\d.,]+).*:",
            ],
            
            # Violation patterns - preserve letter indicators
            "violation_patterns": [
                r"^([a-z]|đ|e|g|h|i|k|l|m|n|o|p|q|r|s|t|u|v|w|x|y|z)\)\s*(.*)",     # a) b) c) đ) e) g) h) etc - preserve letter
                r"^-\s*(.*)",               # - violation  
                r"^\+\s*(.*)",              # + violation
                r"^\*\s*(.*)",              # * violation
            ],
            
            # Additional measure patterns
            "measure_patterns": [
                r"tước quyền sử dụng.*từ\s+(\d+).*đến\s+(\d+)\s+tháng",
                r"tịch thu.*phương tiện",
                r"buộc.*khôi phục",
                r"tạm giữ.*phương tiện"
            ],
            
            # Additional penalty patterns (hình thức phạt bổ sung)
            "additional_penalty_patterns": [
                r"^([a-z]|đ|e|g|h)\)\s*Thực hiện hành vi.*bị\s*(tịch thu|tước quyền|buộc phải).*",
                r"Thực hiện hành vi.*bị\s*(tịch thu|tước quyền|buộc phải).*",
            ],
            
            # Amendment patterns (for updates)
            "amendment_patterns": [
                r"Nghị định\s+(\d+/\d+/NĐ-CP)",
                r"sửa đổi.*bổ sung",
                r"thay thế.*khoản.*điều"
            ]
        }
    
    def calculate_file_hash(self, file_path):
        """Calculate hash of file to detect changes"""
        hasher = hashlib.md5()
        try:
            with open(file_path, 'rb') as f:
                buf = f.read()
                hasher.update(buf)
            return hasher.hexdigest()
        except:
            return None
    
    def detect_document_type(self, file_path):
        """Detect document type and structure"""
        _, ext = os.path.splitext(file_path)
        
        if ext.lower() not in self.config["supported_formats"]:
            raise ValueError(f"Unsupported format: {ext}")
        
        # Try to extract sample content
        sample_content = self.extract_sample_content(file_path)
        
        # Detect structure patterns
        structure = {
            "format": ext.lower(),
            "decree_type": self.detect_decree_type(sample_content),
            "has_amendments": self.detect_amendments(sample_content),
            "language": "vietnamese",
            "confidence": 0.0
        }
        
        return structure
    
    def extract_sample_content(self, file_path):
        """Extract first few lines for structure detection"""
        _, ext = os.path.splitext(file_path)
        
        if ext.lower() == '.docx':
            try:
                doc = Document(file_path)
                return [para.text.strip() for para in doc.paragraphs[:20] if para.text.strip()]
            except:
                return []
        elif ext.lower() == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return [line.strip() for line in f.readlines()[:20] if line.strip()]
            except:
                return []
        
        return []
    
    def detect_decree_type(self, content):
        """Detect the type of decree/regulation"""
        content_text = ' '.join(content).lower()
        
        if "nghị định" in content_text and "100/2019" in content_text:
            return "ND100-2019"
        elif "nghị định" in content_text and "123/2021" in content_text:
            return "ND123-2021"
        elif "nghị định" in content_text and "168/2024" in content_text:
            return "ND168-2024"
        elif "luật giao thông" in content_text:
            return "traffic_law"
        else:
            return "unknown"
    
    def detect_amendments(self, content):
        """Detect if document contains amendments"""
        content_text = ' '.join(content).lower()
        
        for pattern in self.patterns["amendment_patterns"]:
            if re.search(pattern, content_text, re.IGNORECASE):
                return True
        return False
    
    def extract_from_docx(self, file_path):
        """Extract content from DOCX file with enhanced parsing"""
        try:
            doc = Document(file_path)
            
            # Extract all content
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text.strip())
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "text_content": content,
                "tables": tables,
                "extraction_method": "python-docx",
                "timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            raise Exception(f"Failed to extract from DOCX: {e}")
    
    def parse_with_adaptive_patterns(self, content, document_type):
        """Parse content using adaptive patterns based on document type"""
        if not content or not content.get("text_content"):
            return None
        
        text_lines = content["text_content"]
        parsed_articles = {}
        
        # Select appropriate patterns based on document type
        article_patterns = self.patterns["article_patterns"]
        section_patterns = self.patterns["section_patterns"]
        violation_patterns = self.patterns["violation_patterns"]
        
        current_article = None
        current_section = None
        current_violations = []
        current_fine_range = None
        current_additional_measures = []
        current_additional_penalties = []
        
        for i, line in enumerate(text_lines):
            line = line.strip()
            
            # Try each article pattern
            article_match = None
            for pattern in article_patterns:
                article_match = re.match(pattern, line, re.IGNORECASE)
                if article_match:
                    break
            
            if article_match:
                # Save previous article
                if current_article and current_section is not None:
                    self._save_current_section(
                        parsed_articles, current_article, current_section,
                        current_fine_range, current_violations, current_additional_measures
                    )
                
                # Start new article
                article_num = article_match.group(1)
                article_title = article_match.group(2)
                current_article = f"dieu_{article_num}"
                
                if current_article not in parsed_articles:
                    parsed_articles[current_article] = {
                        "title": article_title,
                        "sections": []
                    }
                
                current_section = None
                current_violations = []
                current_fine_range = None
                current_additional_measures = []
                continue
            
            # Try section patterns
            section_match = None
            for pattern in section_patterns:
                section_match = re.match(pattern, line)
                if section_match:
                    break
            
            if section_match:
                # Save previous section
                if current_article and current_section is not None:
                    self._save_current_section(
                        parsed_articles, current_article, current_section,
                        current_fine_range, current_violations, current_additional_measures
                    )
                
                # Start new section
                current_section = section_match.group(1)
                min_fine = section_match.group(2).replace(".", "").replace(",", "")
                max_fine = section_match.group(3).replace(".", "").replace(",", "")
                current_fine_range = f"{min_fine} - {max_fine} VNĐ"
                current_violations = []
                current_additional_measures = []
                current_additional_penalties = []
                continue
            
            # Try violation patterns
            violation_match = None
            for pattern in violation_patterns:
                violation_match = re.match(pattern, line)
                if violation_match:
                    break
            
            if violation_match:
                # Check if it's a lettered violation (a), b), c), etc.)
                if len(violation_match.groups()) > 1 and violation_match.group(1):
                    # Keep the letter indicator for lettered violations
                    letter_indicator = violation_match.group(1)
                    violation_text = violation_match.group(2)
                    violation_text = self._clean_violation_text(violation_text)
                    if violation_text:
                        full_violation = f"{letter_indicator}) \"{violation_text}\""
                        current_violations.append(full_violation)
                else:
                    # For non-lettered violations (-, +, *, etc.)
                    violation_text = violation_match.group(1)
                    violation_text = self._clean_violation_text(violation_text)
                    if violation_text:
                        current_violations.append(violation_text)
                continue
            
            # Check for additional measures
            for pattern in self.patterns["measure_patterns"]:
                if re.search(pattern, line, re.IGNORECASE):
                    measure = self._extract_measure(line, pattern)
                    if measure and measure not in current_additional_measures:
                        current_additional_measures.append(measure)
        
        # Save the last article/section
        if current_article and current_section is not None:
            self._save_current_section(
                parsed_articles, current_article, current_section,
                current_fine_range, current_violations, current_additional_measures
            )
        
        return parsed_articles
    
    def _save_current_section(self, parsed_articles, article_key, section_num, 
                             fine_range, violations, measures, penalties=None):
        """Helper to save current section data"""
        section_data = {
            "section": f"Khoản {section_num}",
            "fine_range": fine_range,
            "violations": violations.copy()
        }
        
        if measures:
            section_data["additional_measures"] = measures.copy()
        
        if penalties:
            section_data["additional_penalties"] = penalties.copy()
        
        parsed_articles[article_key]["sections"].append(section_data)
    
    def _clean_violation_text(self, text):
        """Clean and standardize violation text"""
        if not text:
            return ""
        
        # Remove reference patterns
        text = re.sub(r',\s*trừ.*?;', '', text)
        text = re.sub(r';\s*$', '', text)
        text = re.sub(r'\s*\(.*?\)\s*', ' ', text)  # Remove parenthetical notes
        
        # Clean whitespace
        text = ' '.join(text.split())
        
        # Capitalize first letter
        if text and text[0].islower():
            text = text[0].upper() + text[1:]
        
        return text.strip()
    
    def _extract_measure(self, line, pattern):
        """Extract additional measures from line"""
        if "tước quyền" in line.lower():
            match = re.search(r"từ\s+(\d+).*đến\s+(\d+)\s+tháng", line)
            if match:
                return f"Tước quyền sử dụng giấy phép lái xe từ {match.group(1)} đến {match.group(2)} tháng"
        elif "tịch thu" in line.lower():
            return "Tịch thu phương tiện"
        elif "buộc" in line.lower():
            return "Buộc khôi phục lại tình trạng ban đầu"
        elif "tạm giữ" in line.lower():
            return "Tạm giữ phương tiện"
        
        return None