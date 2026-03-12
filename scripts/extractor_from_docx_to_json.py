#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to extract data from ND168-2024.docx and create nghi_dinh_168_2024.json
with the same format as existing legal documents
"""

import json
import os
import sys
import re
from datetime import datetime
from docx import Document

# Add the project root to Python path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(project_root)

class ND168Extractor:
    """Extractor specifically for Nghị định 168/2024/NĐ-CP"""
    
    def __init__(self):
        self.patterns = self.setup_patterns()
        
    def setup_patterns(self):
        """Setup regex patterns for parsing"""
        return {
            "article_patterns": [
                r"Điều\s+(\d+)\.\s*(.*)",
                r"Điều\s+(\d+):\s*(.*)",
                r"ĐIỀU\s+(\d+)\.\s*(.*)",
            ],
            "section_patterns": [
                r"(\d+)\.\s*Phạt tiền từ\s*([\d.,]+)\s*đồng đến\s*([\d.,]+)\s*đồng",
                r"(\d+)\.\s*Phạt tiền từ\s*([\d.,]+)\s*VNĐ đến\s*([\d.,]+)\s*VNĐ",
                r"Khoản\s+(\d+)\.\s*Phạt tiền từ\s*([\d.,]+)\s*đến\s*([\d.,]+)",
                r"(\d+)\.\s*Phạt tiền.*?từ\s*([\d.,]+).*?đến\s*([\d.,]+)",
                r"(\d+)\.?\s*Phạt tiền từ\s*([\d.,]+)\s*đồng\s*đến\s*([\d.,]+)\s*đồng",
                r"(\d+)\.?\s*Phạt tiền từ\s*([\d.,]+)\s*VNĐ\s*đến\s*([\d.,]+)\s*VNĐ",
            ],
            "violation_patterns": [
                r"^([a-z]|đ)\)\s*(.*)",
                r"^-\s*(.*)",
                r"^\+\s*(.*)",
                r"^\*\s*(.*)",
                r"^\d+\)\s*(.*)",
                r"đối với hành vi.*?(.*)",
                r".*đối với.*?(vi phạm.*)",
            ],
            "measure_patterns": [
                r"tước quyền sử dụng.*từ\s+(\d+).*đến\s+(\d+)\s+tháng",
                r"tịch thu.*phương tiện",
                r"buộc.*khôi phục",
                r"tạm giữ.*phương tiện"
            ]
        }
    
    def extract_from_docx(self, file_path):
        """Extract content from DOCX file"""
        try:
            doc = Document(file_path)
            
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text.strip())
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "text_content": content,
                "tables": tables,
                "extraction_method": "python-docx",
                "timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            raise Exception(f"Failed to extract from DOCX: {e}")
    
    def parse_content(self, content):
        """Parse extracted content into structured format"""
        if not content or not content.get("text_content"):
            return None
        
        text_lines = content["text_content"]
        parsed_articles = {}
        
        current_article = None
        current_section = None
        current_violations = []
        current_fine_range = None
        current_additional_measures = []
        current_article_content = []
        
        for i, line in enumerate(text_lines):
            line = line.strip()
            if not line:
                continue
            
            # Try article patterns
            article_match = None
            for pattern in self.patterns["article_patterns"]:
                article_match = re.match(pattern, line, re.IGNORECASE)
                if article_match:
                    break
            
            if article_match:
                # Save previous article
                if current_article:
                    # Save any pending section
                    if current_section is not None:
                        self._save_current_section(
                            parsed_articles, current_article, current_section,
                            current_fine_range, current_violations, current_additional_measures
                        )
                    # If no sections but has content, save as general content
                    elif current_article_content and not parsed_articles[current_article]["sections"]:
                        parsed_articles[current_article]["content"] = "\n".join(current_article_content)
                
                # Start new article
                article_num = article_match.group(1)
                article_title = article_match.group(2).strip()
                current_article = f"dieu_{article_num}"
                
                parsed_articles[current_article] = {
                    "title": article_title,
                    "sections": [],
                    "article_number": int(article_num)
                }
                
                current_section = None
                current_violations = []
                current_fine_range = None
                current_additional_measures = []
                current_article_content = []
                continue
            
            # If we're in an article, collect content
            if current_article:
                # Try section patterns
                section_match = None
                for pattern in self.patterns["section_patterns"]:
                    section_match = re.search(pattern, line)
                    if section_match:
                        break
                
                if section_match:
                    # Save previous section
                    if current_section is not None:
                        self._save_current_section(
                            parsed_articles, current_article, current_section,
                            current_fine_range, current_violations, current_additional_measures
                        )
                    
                    # Start new section
                    current_section = section_match.group(1)
                    min_fine = self._clean_number(section_match.group(2))
                    max_fine = self._clean_number(section_match.group(3))
                    current_fine_range = f"{min_fine} - {max_fine} VNĐ"
                    current_violations = []
                    current_additional_measures = []
                    continue
                
                # Try violation patterns (only if we have a section)
                if current_section is not None:
                    # Special handling for "đối với hành vi" and "đối với người" format
                    if "đối với hành vi" in line.lower():
                        # Extract the violation description after "đối với hành vi"
                        violation_match = re.search(r"đối với hành vi\s*(.*)", line, re.IGNORECASE)
                        if violation_match:
                            violation_text = self._clean_violation_text(violation_match.group(1))
                            if violation_text and len(violation_text) > 10:
                                current_violations.append(violation_text)
                            continue
                    elif "đối với người" in line.lower():
                        # Extract the violation description after "đối với người"
                        violation_match = re.search(r"đối với người\s*(.*)", line, re.IGNORECASE)
                        if violation_match:
                            violation_text = self._clean_violation_text(violation_match.group(1))
                            if violation_text and len(violation_text) > 10:
                                current_violations.append(violation_text)
                            continue
                    
                    # Regular violation patterns
                    violation_match = None
                    for pattern in self.patterns["violation_patterns"]:
                        violation_match = re.match(pattern, line)
                        if violation_match:
                            break
                    
                    if violation_match:
                        violation_text = violation_match.group(2) if len(violation_match.groups()) > 1 else violation_match.group(1)
                        violation_text = self._clean_violation_text(violation_text)
                        
                        if violation_text and len(violation_text) > 10:  # Filter out very short entries
                            current_violations.append(violation_text)
                        continue
                    
                    # Check for additional measures
                    for pattern in self.patterns["measure_patterns"]:
                        if re.search(pattern, line, re.IGNORECASE):
                            measure = self._extract_measure(line, pattern)
                            if measure and measure not in current_additional_measures:
                                current_additional_measures.append(measure)
                
                # Collect general article content (for articles without sections)
                if current_section is None and len(line) > 20:  # Only meaningful content
                    current_article_content.append(line)
        
        # Save the last article/section
        if current_article:
            if current_section is not None:
                self._save_current_section(
                    parsed_articles, current_article, current_section,
                    current_fine_range, current_violations, current_additional_measures
                )
            elif current_article_content and not parsed_articles[current_article]["sections"]:
                parsed_articles[current_article]["content"] = "\n".join(current_article_content)
        
        return parsed_articles
    
    def _save_current_section(self, parsed_articles, article_key, section_num, 
                             fine_range, violations, measures):
        """Save current section data"""
        # Always save if we have a fine_range (indicating a valid penalty section)
        if not fine_range:
            return
            
        section_data = {
            "section": f"Khoản {section_num}",
            "fine_range": fine_range,
            "violations": violations.copy() if violations else ["Nội dung vi phạm không được chỉ định cụ thể"]
        }
        
        if measures:
            section_data["additional_measures"] = measures.copy()
        
        parsed_articles[article_key]["sections"].append(section_data)
    
    def _clean_number(self, number_str):
        """Clean number string, removing dots and commas"""
        if not number_str:
            return "0"
        return number_str.replace(".", "").replace(",", "")
    
    def _clean_violation_text(self, text):
        """Clean and standardize violation text"""
        if not text:
            return ""
        
        # Remove reference patterns and cleanup
        text = re.sub(r',\s*trừ.*?;', '', text)
        text = re.sub(r';\s*$', '', text)
        text = re.sub(r'\s*\(.*?\)\s*', ' ', text)
        
        # Clean whitespace
        text = ' '.join(text.split())
        
        # Capitalize first letter
        if text and text[0].islower():
            text = text[0].upper() + text[1:]
        
        return text.strip()
    
    def _extract_measure(self, line, pattern):
        """Extract additional measures from line"""
        if "tước quyền" in line.lower():
            match = re.search(r"từ\s+(\d+).*đến\s+(\d+)\s+tháng", line)
            if match:
                return f"Tước quyền sử dụng giấy phép lái xe từ {match.group(1)} đến {match.group(2)} tháng"
        elif "tịch thu" in line.lower():
            return "Tịch thu phương tiện"
        elif "buộc" in line.lower():
            return "Buộc khôi phục lại tình trạng ban đầu"
        elif "tạm giữ" in line.lower():
            return "Tạm giữ phương tiện"
        
        return None
    
    def create_document_structure(self, parsed_articles):
        """Create the complete document structure similar to existing files"""
        
        # Count total articles and chapters
        total_articles = len(parsed_articles)
        
        # Create chapters structure (this would need to be customized based on actual content)
        chapters = [
            {
                "chapter": 1,
                "title": "Những quy định chung",
                "articles": [f"Điều {i}" for i in range(1, 5)]
            },
            {
                "chapter": 2,
                "title": "Vi phạm hành chính trong lĩnh vực giao thông đường bộ",
                "articles": [f"Điều {i}" for i in range(5, 19)]
            },
            {
                "chapter": 3,
                "title": "Vi phạm hành chính trong lĩnh vực giao thông đường sắt",
                "articles": [f"Điều {i}" for i in range(19, 22)]
            },
            {
                "chapter": 4,
                "title": "Thẩm quyền xử phạt và áp dụng biện pháp khắc phục hậu quả",
                "articles": [f"Điều {i}" for i in range(22, 28)]
            },
            {
                "chapter": 5,
                "title": "Điều khoản thi hành",
                "articles": [f"Điều {i}" for i in range(28, total_articles + 1)]
            }
        ]
        
        # Create the main structure
        document_structure = {
            "document_info": {
                "title": "Nghị định 168/2024/NĐ-CP",
                "full_title": "Nghị định về sửa đổi, bổ sung một số điều của Nghị định số 100/2019/NĐ-CP về xử phạt vi phạm hành chính trong lĩnh vực giao thông đường bộ và đường sắt",
                "issued_date": "2024-12-27",
                "effective_date": "2025-01-01",
                "issued_by": "Chính phủ",
                "amendments": [],
                "total_articles": total_articles,
                "total_chapters": len(chapters),
                "description": f"Nghị định sửa đổi, bổ sung Nghị định 100/2019/NĐ-CP - {total_articles} điều với nội dung được trích xuất tự động",
                "last_updated": datetime.now().strftime("%Y-%m-%d"),
                "update_source": "Automated extraction from ND168-2024.docx"
            },
            "structure": {
                "chapters": chapters
            },
            "key_articles": {}
        }
        
        # Add parsed articles to key_articles
        for article_key, article_data in parsed_articles.items():
            # Include all articles that have either sections or content
            if article_data.get("sections") or article_data.get("content"):
                article_entry = {
                    "title": article_data["title"],
                    "source_document": "ND168-2024.docx",
                    "extraction_date": datetime.now().strftime("%Y-%m-%d")
                }
                
                # Add sections if available
                if article_data.get("sections"):
                    article_entry["sections"] = article_data["sections"]
                
                # Add general content if available (for articles without fine structure)
                if article_data.get("content"):
                    article_entry["content"] = article_data["content"]
                
                document_structure["key_articles"][article_key] = article_entry
        
        return document_structure
    
    def process_document(self, input_path, output_path):
        """Main processing function"""
        print(f"Processing document: {input_path}")
        
        # Extract content from DOCX
        print("Extracting content from DOCX...")
        content = self.extract_from_docx(input_path)
        print(f"Extracted {len(content['text_content'])} paragraphs")
        
        # Parse content
        print("Parsing content...")
        parsed_articles = self.parse_content(content)
        print(f"Parsed {len(parsed_articles)} articles")
        
        # Create document structure
        print("Creating document structure...")
        document_structure = self.create_document_structure(parsed_articles)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Save to JSON
        print(f"Saving to: {output_path}")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(document_structure, f, ensure_ascii=False, indent=2)
        
        print("Extraction completed successfully!")
        print(f"Total articles extracted: {len(document_structure['key_articles'])}")
        
        return document_structure

def main():
    """Main function"""
    # Paths
    input_path = os.path.join(project_root, "docs", "ND168-2024.docx")
    output_path = os.path.join(project_root, "data", "raw", "legal_documents", "nghi_dinh_168_2024.json")
    
    # Check if input file exists
    if not os.path.exists(input_path):
        print(f"Error: Input file not found: {input_path}")
        return
    
    # Create extractor and process
    extractor = ND168Extractor()
    try:
        document_structure = extractor.process_document(input_path, output_path)
        print("\n" + "="*50)
        print("EXTRACTION SUMMARY")
        print("="*50)
        print(f"Document: {document_structure['document_info']['title']}")
        print(f"Total articles: {len(document_structure['key_articles'])}")
        print(f"Output file: {output_path}")
        print(f"File size: {os.path.getsize(output_path)} bytes")
        print("="*50)
        
    except Exception as e:
        print(f"Error during extraction: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()